import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import PyPDF2

# Trigger word list (as a raw string for readability)
trigger_words_raw = """
accessible
activism
activists
advocacy
advocate
advocates
affirming care
all-inclusive
allyship
anti-racism
antiracist
assigned at birth
assigned female at birth
assigned male at birth
at risk
barrier
barriers
belong
bias
biased
biased toward
biases
biases towards
biologically female
biologically male
BIPOC
Black
breastfeed + people
breastfeed + person
chestfeed + people
chestfeed + person
clean energy
climate crisis
climate science
commercial sex worker
community diversity
community equity
confirmation bias
cultural competence
cultural differences
cultural heritage
cultural sensitivity
culturally appropriate
culturally responsive
DEI
DEIA
DEIAB
DEIJ
disabilities
disability
discriminated
discrimination
discriminatory
disparity
disparities
diverse
diverse backgrounds
diverse communities
diverse community
diverse group
diverse groups
diversified
diversify
diversifying
diversity
enhance the diversity
enhancing diversity
environmental quality
equal opportunity
equality
equitable
equitableness
equity
ethnicity
excluded
exclusion
expression
female
females
feminism
fostering inclusivity
GBV
gender
gender based
gender based violence
gender diversity
gender identity
gender ideology
gender-affirming care
genders
Gulf of Mexico
hate speech
health disparity
health equity
hispanic minority
historically
identity
immigrants
implicit bias
implicit biases
inclusion
inclusive
inclusive leadership
inclusiveness
inclusivity
increase diversity
increase the diversity
indigenous community
inequalities
inequality
inequitable
inequities
inequity
injustice
institutional
intersectional
intersectionality
key groups
key people
key populations
Latinx
LGBT
LGBTQ
marginalize
marginalized
men who have sex with men
mental health
minorities
minority
most risk
MSM
multicultural
Mx
Native American
non-binary
nonbinary
oppression
oppressive
orientation
people + uterus
people-centered care
person-centered
person-centered care
polarization
political
pollution
pregnant people
pregnant person
pregnant persons
prejudice
privilege
privileges
promote diversity
promoting diversity
pronoun
pronouns
prostitute
race
race and ethnicity
racial
racial diversity
racial identity
racial inequality
racial justice
racially
racism
segregation
sense of belonging
sex
sexual preferences
sexuality
social justice
sociocultural
socioeconomic
status
stereotype
stereotypes
systemic
systemically
they/them
trans
transgender
transsexual
trauma
traumatic
tribal
unconscious bias
underappreciated
underprivileged
underrepresentation
underrepresented
underserved
undervalued
victim
victims
vulnerable populations
women
women and underrepresented
"""

# Process raw list into a usable Python list
trigger_words = [line.strip().lower() for line in trigger_words_raw.splitlines() if line.strip()]

# Normalize special cases with "+" for multi-word match
trigger_words = [w.replace(" + ", r"\s+") for w in trigger_words]

# Build regex
trigger_pattern = re.compile(r'\b(?:' + '|'.join(trigger_words) + r')\b', re.IGNORECASE)

# GUI file picker
def choose_file():
    root = tk.Tk()
    root.withdraw()
    filepath = filedialog.askopenfilename(
        title="Select a document to scan",
        filetypes=[("All supported", "*.txt *.docx *.pdf"), ("Text files", "*.txt"), ("Word documents", "*.docx"), ("PDF files", "*.pdf")]
    )
    return filepath

# Extract text from file
def extract_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.txt':
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        text = ""
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + '\n'
        return text
    elif ext == '.docx':
        doc = Document(filepath)
        return '\n'.join([p.text for p in doc.paragraphs])
    else:
        raise ValueError("Unsupported file type")

# Create highlighted Word doc
def create_highlighted_doc(original_text, matches, output_path):
    doc = Document()
    for line in original_text.splitlines():
        paragraph = doc.add_paragraph()
        start = 0
        for match in trigger_pattern.finditer(line):
            # Add text before match
            paragraph.add_run(line[start:match.start()])
            # Add highlighted match
            run = paragraph.add_run(match.group())
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            run.bold = True
            start = match.end()
        # Add remainder of line
        paragraph.add_run(line[start:])
    doc.save(output_path)

# Main function
def main():
    filepath = choose_file()
    if not filepath:
        messagebox.showinfo("Canceled", "No file selected.")
        return

    text = extract_text(filepath)
    matches = trigger_pattern.findall(text)
    match_count = len(matches)

    # Save new docx with highlights
    output_path = os.path.splitext(filepath)[0] + "_highlighted.docx"
    create_highlighted_doc(text, matches, output_path)

    print(f"\n✅ Scan complete. Found {match_count} trigger word{'s' if match_count != 1 else ''}.")
    print(f"📄 Highlighted version saved to: {output_path}")

if __name__ == "__main__":
    main()